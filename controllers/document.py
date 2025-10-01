from pathlib import Path
from docx import Document as DocxDocument
from copy import deepcopy

class Document:
    def create_act(self, context: dict, month: int):
        doc = DocxDocument('./assets/Акт.docx')

        # Заменяем плейсхолдеры в абзацах
        for p in doc.paragraphs:
            self._replace_in_paragraph(p, context)

        # Заменяем плейсхолдеры в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph(p, context)

        # Заполняем таблицу продуктов
        products = context.get("products", [])
        full_hours = context.get("full_hours", "")
        self._fill_products_table(doc, products, full_hours)

        # Создаём папку для сохранения, если её нет
        save_path = Path(f'./acts/Акт за {month}-ый месяц.docx')
        save_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(save_path)

        return save_path

    @staticmethod
    def _run_has_picture(run):
        """
        Проверяет, содержит ли run картинку.
        """
        for pic in run.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic'):
            if pic is not None:
                return True
        return False

    def _replace_in_paragraph(self, paragraph, context: dict):
        """
        Заменяет плейсхолдеры вида {{key}} на значения из context,
        сохраняя форматирование и картинки.
        Работает даже если плейсхолдер разбит на несколько run.
        """
        # Собираем все текстовые run (картинки пропускаем)
        text_runs = [r for r in paragraph.runs if not self._run_has_picture(r)]
        if not text_runs:
            return

        # Собираем полный текст
        full_text = "".join(r.text for r in text_runs)

        # Заменяем все плейсхолдеры
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"
            full_text = full_text.replace(placeholder, str(value))

        # Размещаем текст обратно по run, не трогая картинки
        idx = 0
        for r in text_runs:
            length = len(r.text)
            r.text = full_text[idx:idx+length]
            idx += length

    def _fill_products_table(self, doc, products, full_hours):
        if not products:
            return

        table = doc.tables[0]
        template_row = table.rows[1]  # строка-шаблон продукта

        # Удаляем все строки кроме заголовка
        for i in range(len(table.rows) - 1, 0, -1):
            table._tbl.remove(table.rows[i]._tr)

        # Добавляем строки продуктов
        for product in products:
            new_row = deepcopy(template_row)
            for idx, cell in enumerate(new_row.cells):
                if idx == 0:
                    cell.text = str(product["product_index"])
                elif idx == 1:
                    cell.text = product["product"]
                elif idx == 2:
                    cell.text = str(product["product_hours"])
            table._tbl.append(new_row._tr)

        # Итоговая строка (копируем шаблон, чтобы сохранить ширину ячеек)
        total_row = deepcopy(template_row)
        total_row.cells[0].text = ""
        total_row.cells[1].text = "Итого"
        total_row.cells[2].text = str(full_hours)
        table._tbl.append(total_row._tr)
